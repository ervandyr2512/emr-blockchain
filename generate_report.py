#!/usr/bin/env python3
"""
Generate FINAL updated EMR Blockchain project report as .docx
Reflects all features as of latest commit.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size=11, color=None, font_name="Times New Roman"):
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.name = font_name
    if color: run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_heading(text, level=level); p.alignment = align
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0x1a,0x37,0x6c) if level==1 else RGBColor(0x1e,0x5a,0x9c)
    return p

def subheading(text, level=2): return heading(text, level=level)

def para(text="", bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None):
    p = doc.add_paragraph(); p.alignment = align
    r = p.add_run(text); set_run_font(r, bold=bold, italic=italic, size=size, color=color)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix); set_run_font(r1, bold=True, size=11)
        r2 = p.add_run(text); set_run_font(r2, size=11)
    else:
        r = p.add_run(text); set_run_font(r, size=11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    shd = OxmlElement("w:pPr"); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),"F5F5F5")
    shd.append(sh); p._p.insert(0, shd)
    r = p.add_run(text); r.font.name = "Courier New"; r.font.size = Pt(9)
    return p

def styled_table(headers, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"1A376C"); tcPr.append(shd)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.name = "Times New Roman"
        r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    return t

def add_row(t, cells, bold_first=False, bg=None):
    row = t.add_row()
    for i, text in enumerate(cells):
        c = row.cells[i]; c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]; r = p.add_run(str(text))
        r.bold = bold_first and i==0; r.font.name = "Times New Roman"; r.font.size = Pt(10)
        if bg:
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
            shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),bg); tcPr.append(shd)
    return row

def divider():
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("─"*80); r.font.name="Courier New"; r.font.size=Pt(8)
    r.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)

def spacer(n=1):
    for _ in range(n): doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

spacer(3)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN TUGAS AKHIR")
set_run_font(r, bold=True, size=20, color=(0x1a,0x37,0x6c))
spacer()

for txt, sz in [("SISTEM ELECTRONIC MEDICAL RECORD (EMR)",18),
                ("BERBASIS TEKNOLOGI BLOCKCHAIN",18)]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(txt), bold=True, size=sz, color=(0x1a,0x37,0x6c))

spacer(2); divider(); spacer()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("Disusun oleh:"), size=12)
spacer()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("Ervandy Rangganata"), bold=True, size=14, color=(0x1a,0x37,0x6c))

spacer(3); divider(); spacer(2)

t = doc.add_table(rows=5, cols=2)
t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
link_data = [
    ("GitHub Repository",        "https://github.com/ervandyr2512/emr-blockchain"),
    ("Aplikasi Web (Vercel)",     "https://emr-blockchain.vercel.app"),
    ("Smart Contract (Sepolia)",  "https://sepolia.etherscan.io/address/0xBdc34e6C796235A1399f2E18e350e62E24b1Dca5"),
    ("Firebase Console",         "https://console.firebase.google.com/project/emr-blockchain"),
    ("Firebase Database",        "https://emr-blockchain-default-rtdb.asia-southeast1.firebasedatabase.app"),
]
for i,(label,url) in enumerate(link_data):
    row = t.rows[i]
    r0 = row.cells[0].paragraphs[0].add_run(label)
    r0.bold=True; r0.font.name="Times New Roman"; r0.font.size=Pt(10)
    r1 = row.cells[1].paragraphs[0].add_run(url)
    r1.font.name="Courier New"; r1.font.size=Pt(9)
    r1.font.color.rgb = RGBColor(0x00,0x56,0xB3)
    tc = row.cells[0]._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"EEF2FF"); tcPr.append(shd)

spacer(3)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("2024 / 2025"), size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DAFTAR ISI
# ══════════════════════════════════════════════════════════════════════════════

heading("DAFTAR ISI", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)

toc = [
    ("I.",    "Pendahuluan",                               "3"),
    ("II.",   "Deskripsi Proyek",                          "4"),
    ("III.",  "Arsitektur Sistem",                         "5"),
    ("IV.",   "Smart Contract — EMRv2.sol",                "7"),
    ("V.",    "Dokumentasi Fungsi CRUD (18 Fungsi)",       "9"),
    ("VI.",   "Dokumentasi Modifier (4 Modifier)",         "18"),
    ("VII.",  "Dokumentasi Event (8 Event)",               "19"),
    ("VIII.", "Sistem Peran (Role-Based Access Control)",  "21"),
    ("IX.",   "Alur Pendaftaran Dokter & Persetujuan Admin","22"),
    ("X.",    "Alur Integrasi Blockchain",                 "24"),
    ("XI.",   "Firebase (Off-Chain Storage)",              "25"),
    ("XII.",  "Struktur Database Firebase",                "26"),
    ("XIII.", "Tampilan Antarmuka Aplikasi",               "27"),
    ("XIV.",  "Tautan Deployment",                         "29"),
    ("XV.",   "Pemenuhan Syarat Tugas",                    "30"),
    ("XVI.",  "Kesimpulan",                                "31"),
]
for num, title, page in toc:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    dots = "." * max(1, 60 - len(num) - len(title))
    set_run_font(p.add_run(f"{num}  {title} {dots} {page}"), size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# I. PENDAHULUAN
# ══════════════════════════════════════════════════════════════════════════════

heading("I. PENDAHULUAN")
para("Rekam medis elektronik (Electronic Medical Record / EMR) merupakan komponen kritis dalam sistem informasi kesehatan modern. Setiap data klinis pasien — mulai dari biodata, hasil pemeriksaan, catatan dokter, hingga resep obat — harus tersimpan secara aman, tidak dapat dimanipulasi, dan dapat diaudit setiap saat oleh pihak yang berwenang.")
para("Proyek ini mengimplementasikan sistem EMR yang menggabungkan dua teknologi: (1) Firebase Realtime Database sebagai penyimpanan data klinis utama (off-chain), dan (2) Smart contract Ethereum pada jaringan Sepolia Testnet sebagai lapisan audit trail yang tidak dapat diubah (on-chain).")

subheading("A. Latar Belakang")
para("Tantangan utama dalam pengelolaan rekam medis konvensional mencakup: risiko pemalsuan data, ketidakjelasan siapa yang mengubah data, serta sulitnya melacak riwayat lengkap perubahan. Teknologi blockchain menawarkan solusi melalui sifat-sifatnya yang immutable, transparent, dan decentralized.")

subheading("B. Tujuan")
for b in [
    "Membangun sistem EMR multi-peran (pasien, dokter, perawat, admin, apoteker).",
    "Mencatat setiap tindakan klinis sebagai transaksi permanen di Ethereum Sepolia.",
    "Menyimpan hash SHA-256 data klinis on-chain sebagai bukti integritas.",
    "Menyediakan alur pendaftaran dokter mandiri dengan mekanisme persetujuan admin.",
    "Memenuhi standar CRUD minimal dengan 18 fungsi smart contract.",
    "Menyediakan antarmuka web yang user-friendly dengan blockchain trail per entri.",
]:
    bullet(b)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# II. DESKRIPSI PROYEK
# ══════════════════════════════════════════════════════════════════════════════

heading("II. DESKRIPSI PROYEK")
subheading("A. Identitas Proyek")
t = styled_table(["Atribut","Detail"], col_widths=[2.0,4.5])
for row in [
    ("Nama Proyek",          "EMR Blockchain — Sistem Rekam Medis Elektronik Berbasis Blockchain"),
    ("Pengembang",           "Ervandy Rangganata"),
    ("Bahasa Pemrograman",   "TypeScript (Next.js 15 App Router), Solidity ^0.8.19"),
    ("Framework Frontend",   "Next.js 15, Tailwind CSS"),
    ("Smart Contract",       "Solidity — EMRv2.sol (file final & production)"),
    ("Blockchain Network",   "Ethereum Sepolia Testnet (Chain ID: 11155111)"),
    ("Off-Chain Database",   "Firebase Realtime Database (asia-southeast1)"),
    ("Autentikasi",          "Firebase Authentication (Email/Password)"),
    ("Deployment",           "Vercel (auto-deploy dari GitHub main branch)"),
    ("Contract Address",     "0xBdc34e6C796235A1399f2E18e350e62E24b1Dca5"),
    ("Tahun",                "2024 / 2025"),
]:
    add_row(t, row, bold_first=True)

spacer()
subheading("B. Teknologi yang Digunakan")
t2 = styled_table(["Teknologi","Fungsi"], col_widths=[2.0,4.5])
for row in [
    ("Next.js 15 App Router", "Framework React SSR/CSR, routing, layout"),
    ("TypeScript",            "Static typing untuk keandalan kode frontend"),
    ("Solidity ^0.8.19",      "Bahasa smart contract Ethereum"),
    ("Hardhat",               "Kompilasi, testing, dan deployment smart contract"),
    ("ethers.js v6",          "Interaksi blockchain dari frontend (connect wallet, sign tx)"),
    ("Firebase Auth",         "Autentikasi email/password multi-role"),
    ("Firebase RTDB",         "Penyimpanan data klinis real-time (off-chain)"),
    ("MetaMask",              "Browser wallet untuk signing transaksi"),
    ("Alchemy",               "RPC Provider Ethereum Sepolia"),
    ("Vercel",                "Hosting + CI/CD otomatis dari GitHub"),
    ("Tailwind CSS",          "Utility-first CSS framework"),
    ("date-fns",              "Formatting tanggal/waktu (dengan guard safeFormat)"),
    ("react-hot-toast",       "Notifikasi toast UI"),
    ("lucide-react",          "Ikon UI"),
    ("SHA-256 (Web Crypto)",  "Hashing payload JSON sebelum disimpan on-chain"),
]:
    add_row(t2, row, bold_first=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# III. ARSITEKTUR SISTEM
# ══════════════════════════════════════════════════════════════════════════════

heading("III. ARSITEKTUR SISTEM")
subheading("A. Gambaran Umum")
para("Sistem menggunakan arsitektur hybrid yang memisahkan penyimpanan data klinis (off-chain) dari lapisan keamanan dan audit trail (on-chain).")
spacer()

code_block("""
┌─────────────────────────────────────────────────────────────────────────┐
│                        PENGGUNA (Browser)                               │
│    Patient │ Doctor │ Nurse │ Admin │ Pharmacist │ Pending Doctor        │
└────────────────────────────┬────────────────────────────────────────────┘
                             │ HTTPS
                             ▼
┌─────────────────────────────────────────────────────────────────────────┐
│               NEXT.JS 15 FRONTEND  (Vercel)                             │
│         https://emr-blockchain.vercel.app                               │
│                                                                         │
│  ┌─────────────┐  ┌──────────────┐  ┌─────────────────────────────┐    │
│  │ Firebase    │  │ Firebase     │  │   Blockchain Layer          │    │
│  │   Auth      │  │   RTDB       │  │  src/lib/blockchain.ts      │    │
│  │ (6 roles)   │  │ (data klinis)│  │  ethers.js v6 + MetaMask    │    │
│  └─────────────┘  └──────────────┘  └─────────────────────────────┘    │
└────────┬─────────────────┬──────────────────────┬───────────────────────┘
         │                 │                      │
         ▼                 ▼                      ▼
┌────────────────┐  ┌──────────────────┐  ┌──────────────────────────────┐
│  Firebase      │  │  Firebase RTDB   │  │  Ethereum Sepolia Testnet    │
│  Authentication│  │ (asia-southeast1)│  │  Smart Contract: EMRv2       │
│                │  │                  │  │  0xBdc34e6C796235A1399f2E    │
│  - Email/pass  │  │  - patients/     │  │  18e350e62E24b1Dca5          │
│  - 6 peran     │  │  - soap_notes/   │  │                              │
│  (termasuk     │  │  - doctor_notes/ │  │  18 fungsi external          │
│  pending_      │  │  - prescriptions/│  │  4 modifier                  │
│  doctor)       │  │  - notifications/│  │  8 event                     │
│                │  │  - doctor_       │  │                              │
│                │  │    applications/ │  │                              │
└────────────────┘  └──────────────────┘  └──────────────────────────────┘
""")

spacer()
subheading("B. Alur Data 8 Langkah")
for step, desc in [
    ("1. Input",         "Pengguna mengisi form klinis di antarmuka web."),
    ("2. Firebase Save", "Data JSON lengkap disimpan ke Firebase RTDB."),
    ("3. SHA-256 Hash",  "sha256() menghitung fingerprint kriptografis dari seluruh payload."),
    ("4. TX Request",    "ethers.js v6 memanggil fungsi contract yang sesuai dengan emrId + dataHash."),
    ("5. MetaMask Sign", "Pengguna menandatangani transaksi melalui MetaMask."),
    ("6. On-Chain",      "Transaksi dikonfirmasi di Ethereum Sepolia; txHash dikembalikan."),
    ("7. Trail Save",    "txHash disimpan ke Firebase RTDB (blockchainTxHash + blockchainHistory)."),
    ("8. Notifikasi",    "Sistem membuat notifikasi real-time untuk peran terkait."),
]:
    bullet(f"{step}: {desc}")

spacer()
subheading("C. Struktur Direktori")
code_block("""emr-blockchain/
├── contracts/
│   ├── EMRv2.sol                  ← Smart contract FINAL (production)
│   └── EMR.sol                    ← Versi 1 (legacy, tidak dipakai)
├── scripts/
│   └── deployV2.js                ← Deploy script EMRv2
├── src/
│   ├── app/
│   │   ├── (auth)/
│   │   │   ├── login/             ← Login (logo → klik kembali ke '/')
│   │   │   ├── register/          ← Pendaftaran pasien (3 langkah)
│   │   │   └── register-dokter/   ← Pendaftaran dokter (3 langkah) ← BARU
│   │   ├── (dashboard)/
│   │   │   ├── admin/             ← Dashboard, pasien, manajemen staff
│   │   │   │   └── staff/         ← Buat akun + tab Permohonan Dokter ← UPDATED
│   │   │   ├── doctor/            ← Dashboard, pasien, EMR, records
│   │   │   ├── nurse/             ← Dashboard, pasien, SOAP, records
│   │   │   ├── pharmacist/        ← Dashboard, resep masuk, riwayat
│   │   │   └── patient/           ← Dashboard, rekam medis
│   │   ├── pending-approval/      ← Halaman tunggu persetujuan dokter ← BARU
│   │   ├── dokter/                ← Direktori dokter (landing)
│   │   ├── tentang-kami/          ← Halaman About
│   │   └── ...
│   ├── components/
│   │   ├── layout/                ← Sidebar (6 role), Header, LandingNavbar
│   │   └── ui/                    ← MedicalHistoryCards (safeFormat), dll.
│   ├── lib/
│   │   ├── blockchain.ts          ← Fungsi blockchain (ethers.js)
│   │   ├── emr.ts                 ← Fungsi Firebase RTDB (+ doctor apps)
│   │   ├── auth.ts                ← Firebase Auth (+ pending_doctor role)
│   │   ├── dateUtils.ts           ← safeFormat() shared utility ← BARU
│   │   ├── hash.ts                ← SHA-256
│   │   └── notifications.ts      ← Real-time notifications
│   └── types/
│       └── index.ts               ← Types (+ DoctorApplication, pending_doctor)
└── .env.local                     ← Tidak di-commit ke git
""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# IV. SMART CONTRACT
# ══════════════════════════════════════════════════════════════════════════════

heading("IV. SMART CONTRACT — EMRv2.sol")
subheading("A. Informasi Deployment")
t = styled_table(["Parameter","Nilai"], col_widths=[2.5,4.0])
for row in [
    ("File",             "contracts/EMRv2.sol"),
    ("Versi Solidity",   "^0.8.19"),
    ("Jaringan",         "Ethereum Sepolia Testnet"),
    ("Chain ID",         "11155111 (0xaa36a7)"),
    ("Contract Address", "0xBdc34e6C796235A1399f2E18e350e62E24b1Dca5"),
    ("Etherscan URL",    "https://sepolia.etherscan.io/address/0xBdc34e6C796235A1399f2E18e350e62E24b1Dca5"),
    ("Deploy Tool",      "Hardhat — npx hardhat run scripts/deployV2.js --network sepolia"),
    ("ABI File",         "src/lib/contract/EMRv2.json"),
]:
    add_row(t, row, bold_first=True)

spacer()
subheading("B. Ringkasan Komponen")
t2 = styled_table(["Kategori","Jumlah","Daftar"], col_widths=[1.8,0.8,3.9])
for row in [
    ("CREATE functions", "7",  "registerPatient, submitSOAP, submitDoctorNote, fulfillPrescription, assignDepartment, selfRegister, assignRole"),
    ("READ functions",   "8",  "getEMRActions, getRole, isRegistered, getTotalActions, getUserProfile, getActionCount, getActionByIndex, getLatestAction"),
    ("UPDATE functions", "2",  "updateRecord, updateUserName"),
    ("DELETE functions", "2",  "deactivateRecord (soft-delete), deactivateUser (soft-delete)"),
    ("TOTAL external",   "19", "18 external/public + 1 internal (_recordAction)"),
    ("Modifier",         "4",  "onlyOwner, onlyAdmin, onlyAuthorized, emrExists"),
    ("Event",            "8",  "PatientRegistered, SOAPSubmitted, DoctorNoteSubmitted, PrescriptionCreated, DepartmentAssigned, RecordUpdated, RoleAssigned, RecordDeactivated"),
    ("Enum",             "2",  "ActionType (6 nilai), Role (6 nilai)"),
    ("Struct",           "2",  "MedicalAction, UserProfile"),
]:
    bg = "F0FFF4" if row[0]=="TOTAL external" else None
    add_row(t2, row, bold_first=True, bg=bg)

spacer()
subheading("C. Enumerasi")
para("ActionType:", bold=True)
code_block("""enum ActionType {
    PATIENT_REGISTERED,     // 0
    SOAP_SUBMITTED,         // 1
    DOCTOR_NOTE_SUBMITTED,  // 2
    PRESCRIPTION_CREATED,   // 3
    RECORD_UPDATED,         // 4
    DEPARTMENT_ASSIGNED     // 5
}""")
para("Role:", bold=True)
code_block("""enum Role {
    NONE, PATIENT, DOCTOR, NURSE, ADMIN, PHARMACIST  // 0–5
}""")
subheading("D. Structs")
code_block("""struct MedicalAction {
    uint256    id;          // global auto-increment
    string     emrId;       // Firebase EMR ID
    string     dataHash;    // SHA-256 hex dari JSON off-chain
    ActionType actionType;
    address    submitter;
    uint256    timestamp;
    bool       isActive;    // soft-delete flag
}
struct UserProfile {
    address wallet;
    Role    role;
    string  name;
    bool    active;
}""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# V. DOKUMENTASI FUNGSI CRUD
# ══════════════════════════════════════════════════════════════════════════════

heading("V. DOKUMENTASI FUNGSI CRUD (18 Fungsi)")

# ── CREATE ────────────────────────────────────────────────────────────────────
subheading("A. CREATE Functions (7 Fungsi)")

create_fns = [
    ("C1","assignRole",
     "function assignRole(address user, Role role, string calldata name) external onlyAdmin",
     "onlyAdmin","RoleAssigned(user, role, name)",
     "Menetapkan peran kepada sebuah wallet address. Hanya admin atau owner yang dapat memanggil. Digunakan untuk mendaftarkan dokter, perawat, dan apoteker secara manual.",
     [("user","address","Wallet address penerima peran"),
      ("role","Role","Peran yang ditetapkan (enum Role)"),
      ("name","string","Nama tampilan di blockchain")]),
    ("C2","selfRegister",
     "function selfRegister(string calldata name) external",
     "—","RoleAssigned(msg.sender, Role.PATIENT, name)",
     "Mendaftarkan diri sebagai pasien. Wallet belum boleh memiliki peran.",
     [("name","string","Nama pasien")]),
    ("C3","registerPatient",
     "function registerPatient(string calldata emrId, string calldata dataHash) external onlyAuthorized",
     "onlyAuthorized","PatientRegistered(emrId, dataHash, msg.sender, block.timestamp)",
     "Mendaftarkan EMR ID pasien baru ke blockchain. Setelah ini EMR ID dapat digunakan oleh fungsi lain.",
     [("emrId","string","EMR ID unik (format: EMR-YYYYMMDD-XXXXX)"),
      ("dataHash","string","SHA-256 dari JSON biodata pasien di Firebase")]),
    ("C4","submitSOAP",
     "function submitSOAP(string calldata emrId, string calldata dataHash) external onlyAuthorized emrExists(emrId)",
     "onlyAuthorized, emrExists","SOAPSubmitted(emrId, dataHash, msg.sender, block.timestamp)",
     "Mencatat pengisian SOAP (Subjective/Objective/Assessment/Plan) oleh perawat.",
     [("emrId","string","EMR ID pasien"),
      ("dataHash","string","SHA-256 dari JSON SOAP")]),
    ("C5","submitDoctorNote",
     "function submitDoctorNote(string calldata emrId, string calldata dataHash) external onlyAuthorized emrExists(emrId)",
     "onlyAuthorized, emrExists","DoctorNoteSubmitted(emrId, dataHash, msg.sender, block.timestamp)",
     "Mencatat catatan pemeriksaan, diagnosis, dan rencana terapi dokter.",
     [("emrId","string","EMR ID pasien"),
      ("dataHash","string","SHA-256 dari JSON catatan dokter")]),
    ("C6","fulfillPrescription",
     "function fulfillPrescription(string calldata emrId, string calldata dataHash) external onlyAuthorized emrExists(emrId)",
     "onlyAuthorized, emrExists","PrescriptionCreated(emrId, dataHash, msg.sender, block.timestamp)",
     "Mencatat penyerahan resep oleh apoteker. Menjamin setiap transaksi obat memiliki jejak audit permanen.",
     [("emrId","string","EMR ID pasien"),
      ("dataHash","string","SHA-256 dari JSON resep + data penyerahan")]),
    ("C7","assignDepartment",
     "function assignDepartment(string calldata emrId, string calldata dataHash) external onlyAdmin emrExists(emrId)",
     "onlyAdmin, emrExists","DepartmentAssigned(emrId, dataHash, msg.sender, block.timestamp)",
     "Admin menugaskan pasien ke departemen/poli tertentu.",
     [("emrId","string","EMR ID pasien"),
      ("dataHash","string","SHA-256 dari JSON penugasan departemen")]),
]

for fn in create_fns:
    cid,name,sig,mod,ev,desc,params = fn
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"({cid}) {name}"), bold=True, size=12, color=(0x1A,0x37,0x6C))
    para(desc)
    para("Signature:", bold=True); code_block(sig)
    para("Parameter:", bold=True)
    pt = styled_table(["Parameter","Tipe","Keterangan"], col_widths=[1.2,1.0,4.3])
    for pr in params: add_row(pt, pr, bold_first=True)
    p2=doc.add_paragraph(); r1=p2.add_run("Modifier: "); set_run_font(r1,bold=True,size=11)
    set_run_font(p2.add_run(mod),size=11)
    p3=doc.add_paragraph(); r2=p3.add_run("Event: "); set_run_font(r2,bold=True,size=11)
    set_run_font(p3.add_run(ev),size=11,color=(0x15,0x6B,0x4F))
    spacer()

doc.add_page_break()

# ── READ ──────────────────────────────────────────────────────────────────────
subheading("B. READ Functions (8 Fungsi)")

read_fns = [
    ("R1","getEMRActions(string emrId) → MedicalAction[]",
     "Mengambil seluruh riwayat tindakan klinis untuk EMR ID tertentu (audit trail lengkap)."),
    ("R2","getRole(address user) → Role",
     "Mengembalikan peran enum dari sebuah wallet address."),
    ("R3","isRegistered(string emrId) → bool",
     "Memverifikasi apakah EMR ID sudah terdaftar di blockchain."),
    ("R4","getTotalActions() → uint256",
     "Total jumlah seluruh tindakan klinis yang pernah direkam (semua pasien)."),
    ("R5","getUserProfile(address user) → UserProfile",
     "Mengambil profil lengkap pengguna: wallet, role, name, active."),
    ("R6","getActionCount(string emrId) → uint256",
     "Jumlah tindakan yang tercatat untuk satu EMR ID tertentu."),
    ("R7","getActionByIndex(string emrId, uint256 index) → MedicalAction",
     "Satu entri tindakan berdasarkan indeks berbasis nol. Revert jika indeks melebihi batas."),
    ("R8","getLatestAction(string emrId) → MedicalAction",
     "Entri aktif terbaru. Iterasi mundur, mencari isActive = true."),
]
t = styled_table(["ID","Signature","Deskripsi"], col_widths=[0.5,2.8,3.2])
for cid,sig,desc in read_fns: add_row(t,(cid,sig,desc), bold_first=True)

spacer()
# ── UPDATE ────────────────────────────────────────────────────────────────────
subheading("C. UPDATE Functions (2 Fungsi)")

for cid,name,sig,mod,ev,desc in [
    ("U1","updateRecord",
     "function updateRecord(string calldata emrId, string calldata dataHash, ActionType actionType) external onlyAuthorized emrExists(emrId)",
     "onlyAuthorized, emrExists","RecordUpdated(emrId, dataHash, actionType, msg.sender, block.timestamp)",
     "Mencatat pembaruan rekam medis. Entri lama TIDAK dihapus — sebuah entri baru RECORD_UPDATED ditambahkan. Seluruh riwayat perubahan dapat ditelusuri sepenuhnya."),
    ("U2","updateUserName",
     "function updateUserName(address user, string calldata newName) external",
     "Validasi manual (user sendiri / admin / owner)","RoleAssigned(user, role, newName)",
     "Memperbarui nama tampilan di profil blockchain. Pengguna dapat memperbarui nama sendiri; admin dapat memperbarui siapapun."),
]:
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"({cid}) {name}"), bold=True, size=12, color=(0x1A,0x37,0x6C))
    para(desc)
    para("Signature:", bold=True); code_block(sig)
    p2=doc.add_paragraph(); set_run_font(p2.add_run("Modifier: "),bold=True,size=11)
    set_run_font(p2.add_run(mod),size=11)
    p3=doc.add_paragraph(); set_run_font(p3.add_run("Event: "),bold=True,size=11)
    set_run_font(p3.add_run(ev),size=11,color=(0x15,0x6B,0x4F))
    spacer()

# ── DELETE ────────────────────────────────────────────────────────────────────
subheading("D. DELETE Functions (2 Fungsi — Soft Delete)")
para("Karena blockchain bersifat immutable, data tidak benar-benar dihapus. Flag isActive diset false; data tetap on-chain untuk audit.", italic=True, color=(0x60,0x60,0x60))
spacer()

for cid,name,sig,mod,ev,desc in [
    ("D1","deactivateRecord",
     "function deactivateRecord(string calldata emrId, uint256 actionIndex) external onlyAdmin",
     "onlyAdmin","RecordDeactivated(emrId, actionIndex, msg.sender, block.timestamp)",
     "Menonaktifkan satu entri tindakan (isActive = false). Entri tetap ada di blockchain untuk keperluan audit."),
    ("D2","deactivateUser",
     "function deactivateUser(address user) external onlyAdmin",
     "onlyAdmin","— (tidak ada event)",
     "Menonaktifkan profil pengguna (active = false). Owner tidak dapat dinonaktifkan."),
]:
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"({cid}) {name}"), bold=True, size=12, color=(0x1A,0x37,0x6C))
    para(desc)
    para("Signature:", bold=True); code_block(sig)
    p2=doc.add_paragraph(); set_run_font(p2.add_run("Modifier: "),bold=True,size=11)
    set_run_font(p2.add_run(mod),size=11)
    p3=doc.add_paragraph(); set_run_font(p3.add_run("Event: "),bold=True,size=11)
    set_run_font(p3.add_run(ev),size=11,color=(0x15,0x6B,0x4F))
    spacer()

subheading("E. Internal Helper")
para("_recordAction(emrId, dataHash, actionType) internal", bold=True)
para("Dipanggil oleh semua fungsi CREATE. Menginkremen actionCount global, membuat struct MedicalAction, dan push ke emrActions[emrId]. Memusatkan logika pencatatan (DRY principle).")
code_block("""function _recordAction(string calldata emrId, string calldata dataHash, ActionType actionType) internal {
    uint256 newId = ++actionCount;
    emrActions[emrId].push(MedicalAction({
        id: newId, emrId: emrId, dataHash: dataHash,
        actionType: actionType, submitter: msg.sender,
        timestamp: block.timestamp, isActive: true
    }));
}""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# VI. MODIFIER
# ══════════════════════════════════════════════════════════════════════════════

heading("VI. DOKUMENTASI MODIFIER (4 Modifier)")

modifiers = [
    ("M1","onlyOwner",
     "Membatasi akses hanya untuk wallet deployer kontrak (owner). Baris pertahanan tertinggi.",
     "Digunakan oleh: onlyAdmin (fallback), dan secara implisit di seluruh alur admin.",
     "require(msg.sender == owner, \"EMRv2: caller is not owner\");"),
    ("M2","onlyAdmin",
     "Membatasi akses untuk pengguna berperan ADMIN atau owner. Fleksibel: owner selalu bisa bertindak sebagai admin.",
     "Digunakan oleh: assignRole, assignDepartment, deactivateRecord, deactivateUser",
     "require(userProfiles[msg.sender].role == Role.ADMIN || msg.sender == owner, ...);"),
    ("M3","onlyAuthorized",
     "Membatasi akses untuk semua pengguna terdaftar (role > NONE) atau owner. Mencakup semua 5 peran aktif.",
     "Digunakan oleh: registerPatient, submitSOAP, submitDoctorNote, fulfillPrescription, updateRecord",
     "require(uint8(userProfiles[msg.sender].role) > 0 || msg.sender == owner, ...);"),
    ("M4","emrExists",
     "Memvalidasi bahwa EMR ID sudah terdaftar (ada dalam registeredEMRs). Mencegah aksi pada pasien yang belum terdaftar.",
     "Digunakan oleh: submitSOAP, submitDoctorNote, fulfillPrescription, assignDepartment, updateRecord",
     "require(registeredEMRs[emrId], \"EMRv2: EMR ID not registered\");"),
]

for mid,name,desc,used,code in modifiers:
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"{mid}. {name}"), bold=True, size=12, color=(0x1A,0x37,0x6C))
    para(desc); para(used, italic=True, color=(0x60,0x60,0x60))
    code_block(code); spacer()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# VII. EVENT
# ══════════════════════════════════════════════════════════════════════════════

heading("VII. DOKUMENTASI EVENT (8 Event)")

events = [
    ("E1","PatientRegistered","registerPatient()",
     "Dipancarkan saat EMR ID pasien baru terdaftar di blockchain.",
     "event PatientRegistered(string indexed emrId, string dataHash, address indexed registrar, uint256 timestamp);"),
    ("E2","SOAPSubmitted","submitSOAP()",
     "Dipancarkan saat perawat menyimpan catatan SOAP.",
     "event SOAPSubmitted(string indexed emrId, string dataHash, address indexed nurse, uint256 timestamp);"),
    ("E3","DoctorNoteSubmitted","submitDoctorNote()",
     "Dipancarkan saat dokter menyimpan catatan pemeriksaan.",
     "event DoctorNoteSubmitted(string indexed emrId, string dataHash, address indexed doctor, uint256 timestamp);"),
    ("E4","PrescriptionCreated","fulfillPrescription()",
     "Dipancarkan saat apoteker mencatat penyerahan resep.",
     "event PrescriptionCreated(string indexed emrId, string dataHash, address indexed pharmacist, uint256 timestamp);"),
    ("E5","DepartmentAssigned","assignDepartment()",
     "Dipancarkan saat admin menugaskan pasien ke departemen.",
     "event DepartmentAssigned(string indexed emrId, string dataHash, address indexed admin, uint256 timestamp);"),
    ("E6","RecordUpdated","updateRecord()",
     "Dipancarkan saat terjadi pembaruan rekam medis. Menyertakan actionType untuk informasi jenis record.",
     "event RecordUpdated(string indexed emrId, string dataHash, ActionType actionType, address indexed updater, uint256 timestamp);"),
    ("E7","RoleAssigned","assignRole(), selfRegister(), updateUserName()",
     "Dipancarkan setiap kali peran atau nama pengguna diperbarui. Satu event untuk tiga skenario.",
     "event RoleAssigned(address indexed user, Role role, string name);"),
    ("E8","RecordDeactivated","deactivateRecord()",
     "Dipancarkan saat admin soft-delete satu entri tindakan klinis.",
     "event RecordDeactivated(string indexed emrId, uint256 actionIndex, address indexed deactivatedBy, uint256 timestamp);"),
]

for eid,name,trigger,desc,code in events:
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"{eid}. {name}"), bold=True, size=12, color=(0x15,0x6B,0x4F))
    para(desc)
    p2=doc.add_paragraph(); set_run_font(p2.add_run("Di-emit oleh: "),bold=True,size=11)
    set_run_font(p2.add_run(trigger),size=11)
    code_block(code); spacer()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# VIII. SISTEM PERAN
# ══════════════════════════════════════════════════════════════════════════════

heading("VIII. SISTEM PERAN (ROLE-BASED ACCESS CONTROL)")
para("Sistem mengimplementasikan 6 peran — 5 peran aktif dan 1 peran sementara (pending_doctor) untuk dokter yang sedang menunggu verifikasi.")
spacer()

t = styled_table(["Peran","Kode","Akses Smart Contract","Halaman Aplikasi","Cara Mendaftar"],
                 col_widths=[1.0,0.6,2.0,1.8,1.1])
for row in [
    ("ADMIN","4","onlyAdmin + onlyAuthorized penuh","/admin, /admin/patients, /admin/staff","Dibuat oleh owner/admin"),
    ("DOCTOR","2","onlyAuthorized: submitDoctorNote, updateRecord","/doctor, /doctor/patients, /doctor/emr, /doctor/records","Via /register-dokter (+ persetujuan admin) atau dibuat admin langsung"),
    ("NURSE","3","onlyAuthorized: submitSOAP, updateRecord","/nurse, /nurse/patients, /nurse/soap, /nurse/records","Dibuat oleh admin"),
    ("PHARMACIST","5","onlyAuthorized: fulfillPrescription, updateRecord","/pharmacist, /pharmacist/prescriptions, /pharmacist/dispensed","Dibuat oleh admin"),
    ("PATIENT","1","onlyAuthorized: registerPatient (read-only klinis)","/patient, /patient/emr","Self-register di /register"),
    ("PENDING_DOCTOR","—","Tidak ada akses fungsi klinis","/pending-approval (waiting screen)","Self-register di /register-dokter (belum disetujui)"),
]:
    add_row(t, row, bold_first=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# IX. ALUR PENDAFTARAN DOKTER
# ══════════════════════════════════════════════════════════════════════════════

heading("IX. ALUR PENDAFTARAN DOKTER & PERSETUJUAN ADMIN")
para("Sistem menyediakan alur mandiri bagi dokter untuk mendaftar tanpa harus menghubungi admin terlebih dahulu. Akun baru langsung dibuat namun dikunci dalam status 'pending' hingga admin menyetujuinya.")
spacer()

subheading("A. Alur Pendaftaran Dokter (Self-Registration)")
for step, desc in [
    ("1. Akses /register-dokter",   "Dokter mengakses halaman pendaftaran khusus dari tombol 'Daftar Sebagai Dokter' di halaman /dokter."),
    ("2. Step 1 — Akun",            "Isi nama lengkap (dengan gelar, mis. dr. Budi, Sp.PD), email, password."),
    ("3. Step 2 — Profil Profesional","Isi no. telepon, spesialisasi, Nomor STR, Nomor SIP, RS/Klinik, pendidikan, tahun pengalaman, bio."),
    ("4. Step 3 — Konfirmasi",      "Review semua data + informasi proses review 1×24 jam."),
    ("5. Submit",                   "signUp() membuat Firebase Auth account dengan role='pending_doctor'. Data aplikasi lengkap disimpan ke doctor_applications/{uid} di Firebase RTDB."),
    ("6. Redirect /pending-approval","Halaman status yang menampilkan data permohonan, badge 'Menunggu Review', dan info proses."),
]:
    bullet(f"{step}: {desc}")

spacer()
subheading("B. Alur Review Admin")
for step, desc in [
    ("1. Admin login",    "Admin login ke dashboard /admin/staff."),
    ("2. Tab Permohonan Dokter", "Tab kedua di halaman Manajemen Staff. Badge merah menunjukkan jumlah permohonan pending."),
    ("3. Lihat detail",   "Klik accordion untuk melihat: nama, email, telepon, STR, SIP, RS/klinik, pendidikan, pengalaman, bio."),
    ("4a. Setujui",       "Klik 'Setujui' → approveDoctorApplication(): update role 'pending_doctor' → 'doctor' di Firebase, status aplikasi → 'approved'."),
    ("4b. Tolak",         "Klik 'Tolak' → modal input alasan → rejectDoctorApplication(): status → 'rejected' + alasan tersimpan."),
    ("5. Dokter login",   "Setelah disetujui: login berhasil dan diarahkan ke /doctor dashboard. Jika ditolak: melihat alasan di /pending-approval."),
]:
    bullet(f"{step}: {desc}")

spacer()
subheading("C. Firebase Functions untuk Doctor Applications")
code_block("""// src/lib/emr.ts
saveDoctorApplication(application: DoctorApplication)  → void
getDoctorApplications()                                 → DoctorApplication[]
getDoctorApplication(uid: string)                       → DoctorApplication | null
approveDoctorApplication(uid, adminName)                → void  // sets role="doctor"
rejectDoctorApplication(uid, adminName, reason)         → void  // sets status="rejected"

// Firebase RTDB path: doctor_applications/{uid}""")

spacer()
subheading("D. DoctorApplication Interface")
code_block("""interface DoctorApplication {
  uid:            string;   // Firebase Auth UID
  email:          string;
  name:           string;   // Nama lengkap dengan gelar
  phone:          string;
  specialization: string;
  strNumber:      string;   // Surat Tanda Registrasi
  sipNumber:      string;   // Surat Izin Praktik
  hospital:       string;
  education:      string;
  experience:     number;   // Tahun pengalaman
  bio:            string;
  status:         "pending" | "approved" | "rejected";
  submittedAt:    string;
  reviewedAt?:    string;
  reviewedBy?:    string;
  rejectReason?:  string;
}""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# X. ALUR INTEGRASI BLOCKCHAIN
# ══════════════════════════════════════════════════════════════════════════════

heading("X. ALUR INTEGRASI BLOCKCHAIN")

subheading("A. Prinsip Hash-Before-Store")
para("Data medis tidak disimpan di blockchain karena privasi dan biaya gas. Sistem menyimpan data di Firebase, kemudian menyimpan SHA-256 hash-nya on-chain sebagai 'fingerprint' integritas.")
code_block("""// src/lib/hash.ts
export async function sha256(data: unknown): Promise<string> {
  const json   = JSON.stringify(data);
  const buffer = new TextEncoder().encode(json);
  const digest = await crypto.subtle.digest("SHA-256", buffer);
  return Array.from(new Uint8Array(digest))
    .map(b => b.toString(16).padStart(2, "0")).join("");
}""")

spacer()
subheading("B. Blockchain Trail Per Entri")
para("Setiap catatan SOAP dan catatan dokter memiliki blockchainHistory subcollection di Firebase, yang menyimpan semua txHash secara kronologis (created / updated):")
code_block("""// Firebase path: soap_notes/{emrId}/{noteId}/blockchainHistory/{pushKey}
interface BlockchainTrailEntry {
  txHash:    string;
  timestamp: string;   // ISO string
  action:    "created" | "updated";
  actorName: string;
}

// Fungsi di emr.ts:
addSOAPBlockchainTrail(emrId, noteId, entry: BlockchainTrailEntry)
addDoctorNoteBlockchainTrail(emrId, noteId, entry: BlockchainTrailEntry)""")

spacer()
subheading("C. Fungsi Blockchain Frontend (src/lib/blockchain.ts)")
t = styled_table(["Fungsi","Memanggil Contract","Keterangan"], col_widths=[2.5,1.8,2.2])
for row in [
    ("blockchainRegisterPatient(emrId, dataHash, cb)",     "registerPatient()",     "Daftarkan EMR baru"),
    ("blockchainSubmitSOAPFull(emrId, dataHash, cb)",      "submitSOAP()",          "Catat SOAP perawat"),
    ("blockchainSubmitDoctorNoteFull(emrId, dataHash, cb)","submitDoctorNote()",    "Catat pemeriksaan dokter"),
    ("blockchainFulfillPrescriptionFull(emrId, dH, cb)",   "fulfillPrescription()", "Catat penyerahan obat"),
    ("blockchainAssignDepartment(emrId, dataHash, cb)",    "assignDepartment()",    "Catat penugasan departemen"),
    ("blockchainUpdateRecord(emrId, dH, actionType, cb)",  "updateRecord()",        "Catat pembaruan rekam medis"),
    ("getEMRActionsFromChain(emrId)",                      "getEMRActions()",       "Baca audit trail dari chain"),
]:
    add_row(t, row, bold_first=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# XI. FIREBASE
# ══════════════════════════════════════════════════════════════════════════════

heading("XI. FIREBASE (OFF-CHAIN STORAGE)")

subheading("A. Konfigurasi")
t = styled_table(["Parameter","Nilai"], col_widths=[2.5,4.0])
for row in [
    ("Project ID",     "emr-blockchain"),
    ("Auth Domain",    "emr-blockchain.firebaseapp.com"),
    ("Database URL",   "https://emr-blockchain-default-rtdb.asia-southeast1.firebasedatabase.app"),
    ("Region",         "asia-southeast1 (Singapura)"),
    ("Auth Methods",   "Email/Password"),
]:
    add_row(t, row, bold_first=True)

spacer()
# ══════════════════════════════════════════════════════════════════════════════
# XII. STRUKTUR DATABASE
# ══════════════════════════════════════════════════════════════════════════════

heading("XII. STRUKTUR DATABASE FIREBASE")

code_block("""Firebase Realtime Database Structure:
│
├── users/{uid}                         ← UserProfile
│   ├── uid, email, name, role, createdAt
│
├── patients/{emrId}                    ← Biodata pasien
│   ├── emrId, uid, firstName, lastName
│   ├── gender, ktpNumber, phone, email
│   ├── address: {street, kelurahan, kecamatan, kota, kodePos}
│   ├── emergencyContact: {name, phone}
│   ├── department, status, blockchainTxHash
│   └── createdAt, updatedAt
│
├── soap_notes/{emrId}/{noteId}         ← Catatan SOAP perawat
│   ├── subjective, objective (VitalSigns), assessment, plan
│   ├── nurseUid, nurseName
│   ├── createdAt, updatedAt
│   ├── blockchainTxHash
│   └── blockchainHistory/{pushKey}     ← Trail per aksi
│       └── {txHash, timestamp, action, actorName}
│
├── doctor_notes/{emrId}/{noteId}       ← Catatan pemeriksaan dokter
│   ├── chiefComplaint, historyPresentIllness, pastMedicalHistory
│   ├── surgicalHistory, medicationHistory, allergy
│   ├── vitalSigns, physicalExamination
│   ├── supportingExams[], workingDiagnosis, differentialDiagnosis
│   ├── managementPlan
│   ├── doctorUid, doctorName
│   ├── createdAt, updatedAt
│   ├── blockchainTxHash
│   └── blockchainHistory/{pushKey}
│
├── prescriptions/{emrId}/{rxId}        ← Resep
│   ├── medications: [{name, dose, frequency, duration, notes}]
│   ├── doctorUid, doctorName
│   ├── pharmacistUid?, pharmacistName?
│   ├── status: pending|processing|dispensed|cancelled
│   ├── createdAt, dispensedAt?
│   └── blockchainTxHash?
│
├── notifications/{pushKey}             ← Notifikasi real-time
│   ├── icon, title, body
│   ├── targetRoles: string[]
│   ├── createdAt, unread
│   └── txHash?, emrId?
│
├── doctor_applications/{uid}           ← Permohonan pendaftaran dokter ← BARU
│   ├── uid, email, name, phone
│   ├── specialization, strNumber, sipNumber
│   ├── hospital, education, experience, bio
│   ├── status: pending|approved|rejected
│   ├── submittedAt
│   └── reviewedAt?, reviewedBy?, rejectReason?
│
└── counters/patients                   ← Auto-increment untuk EMR ID""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# XIII. TAMPILAN ANTARMUKA
# ══════════════════════════════════════════════════════════════════════════════

heading("XIII. TAMPILAN ANTARMUKA APLIKASI")

subheading("A. Halaman Landing (Publik)")
t = styled_table(["URL","Deskripsi"], col_widths=[2.5,4.0])
for row in [
    ("/",                  "Landing page — hero, fitur, CTA daftar"),
    ("/dokter",            "Direktori dokter spesialis + CTA 'Daftar Sebagai Dokter' → /register-dokter"),
    ("/tentang-kami",      "Tim (termasuk dr. Ervandy Rangganata, SpU sebagai Founder), misi, milestone, tech stack"),
    ("/rumah-sakit",       "Daftar rumah sakit mitra"),
    ("/kontak-kami",       "Halaman kontak"),
]:
    add_row(t, row, bold_first=True)

spacer()
subheading("B. Halaman Auth")
t2 = styled_table(["URL","Deskripsi"], col_widths=[2.5,4.0])
for row in [
    ("/login",             "Login — logo klik → kembali ke '/'. Email/password + demo accounts"),
    ("/register",          "Pendaftaran pasien 3 langkah — logo klik → kembali ke '/'"),
    ("/register-dokter",   "Pendaftaran dokter 3 langkah (akun → profil profesional → konfirmasi) ← BARU"),
    ("/pending-approval",  "Halaman status tunggu persetujuan admin untuk pending_doctor ← BARU"),
]:
    add_row(t2, row, bold_first=True)

spacer()
subheading("C. Dashboard per Peran")
t3 = styled_table(["Peran","URL","Fitur Utama"], col_widths=[1.0,2.2,3.3])
for row in [
    ("Admin",      "/admin",                        "Statistik, quick actions, pasien terbaru"),
    ("Admin",      "/admin/patients",               "Daftar semua pasien, filter status, assign departemen + blockchain"),
    ("Admin",      "/admin/patients/[emrId]",       "Detail pasien: biodata, SOAP, catatan dokter, resep"),
    ("Admin",      "/admin/staff",                  "Tab 1: Buat akun staff. Tab 2: Permohonan dokter (approve/reject) ← UPDATED"),
    ("Dokter",     "/doctor",                       "Dashboard: pasien aktif, shortcut"),
    ("Dokter",     "/doctor/patients",              "Pasien dengan tombol 'Periksa' dan 'Rekam Medis'"),
    ("Dokter",     "/doctor/emr/[patientId]",       "Form pemeriksaan + integrasi blockchain"),
    ("Dokter",     "/doctor/records/[patientId]",   "View-only rekam medis lengkap + blockchain trail"),
    ("Perawat",    "/nurse/patients",               "Pasien dengan tombol 'Input SOAP' dan 'Rekam Medis'"),
    ("Perawat",    "/nurse/soap/[patientId]",       "Form SOAP + tanda vital + integrasi blockchain"),
    ("Perawat",    "/nurse/records/[patientId]",    "View-only rekam medis lengkap + blockchain trail"),
    ("Apoteker",   "/pharmacist",                   "Dashboard + antrian resep singkat"),
    ("Apoteker",   "/pharmacist/prescriptions",     "Antrian resep: Lihat (view-only) + Proses (dispense + blockchain)"),
    ("Apoteker",   "/pharmacist/dispensed",         "Riwayat obat diserahkan: accordion + blockchain trail"),
    ("Pasien",     "/patient",                      "Dashboard: ringkasan kondisi, resep terbaru"),
    ("Pasien",     "/patient/emr",                  "Rekam medis sendiri: SOAP, catatan dokter, resep, blockchain trail"),
]:
    add_row(t3, row, bold_first=True)

spacer()
subheading("D. Fitur UI Unggulan")
for f in [
    "Sidebar role-aware: menampilkan hanya menu yang relevan untuk 6 peran (termasuk pending_doctor → /pending-approval).",
    "Logo auth clickable: di halaman login, register, dan register-dokter, logo/judul dapat diklik untuk kembali ke halaman utama '/'.",
    "Blockchain Trail Cards: setiap entri SOAP dan catatan dokter memiliki riwayat txHash lengkap per aksi (dibuat/diperbarui) dalam accordion.",
    "Badge '✏️ Diperbarui': entri yang diedit menampilkan badge + timestamp pembaruan sebagai waktu utama.",
    "Patient Identity Panel di apoteker: panel verifikasi identitas (Nama, EMR ID, No. KTP, Jenis Kelamin, Telepon, Poli) sebelum menyerahkan obat.",
    "safeFormat() utility: mencegah crash 'Invalid Date' ketika field timestamp Firebase kosong/null/undefined.",
    "Doctor Applications tab: admin melihat detail lengkap (STR, SIP, RS, pendidikan) sebelum menyetujui dokter.",
    "Badge merah pada tab Permohonan Dokter: jumlah permohonan pending ditampilkan secara real-time.",
]:
    bullet(f)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# XIV. TAUTAN DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════

heading("XIV. TAUTAN DEPLOYMENT")

t = styled_table(["Platform","Tautan / Informasi"], col_widths=[2.2,4.3])
for row in [
    ("GitHub Repository",        "https://github.com/ervandyr2512/emr-blockchain"),
    ("Aplikasi Web (Vercel)",     "https://emr-blockchain.vercel.app"),
    ("Smart Contract Address",   "0xBdc34e6C796235A1399f2E18e350e62E24b1Dca5"),
    ("Sepolia Etherscan",        "https://sepolia.etherscan.io/address/0xBdc34e6C796235A1399f2E18e350e62E24b1Dca5"),
    ("Blockchain Network",       "Ethereum Sepolia Testnet (Chain ID: 11155111)"),
    ("Firebase Console",         "https://console.firebase.google.com/project/emr-blockchain"),
    ("Firebase Database URL",    "https://emr-blockchain-default-rtdb.asia-southeast1.firebasedatabase.app"),
    ("Firebase Project ID",      "emr-blockchain"),
    ("File Smart Contract",      "contracts/EMRv2.sol"),
    ("ABI Contract",             "src/lib/contract/EMRv2.json"),
]:
    add_row(t, row, bold_first=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# XV. PEMENUHAN SYARAT TUGAS
# ══════════════════════════════════════════════════════════════════════════════

heading("XV. PEMENUHAN SYARAT TUGAS")

subheading("A. Checklist Persyaratan")
t = styled_table(["No.","Syarat","Status","Bukti"], col_widths=[0.4,2.2,0.7,3.2])
for row in [
    ("1",  "Minimal 12 fungsi CRUD",                  "✅ LULUS","18 fungsi (C7+R8+U2+D2), melebihi syarat"),
    ("2",  "Minimal 1 modifier",                       "✅ LULUS","4 modifier: onlyOwner, onlyAdmin, onlyAuthorized, emrExists"),
    ("3",  "Minimal 1 event",                          "✅ LULUS","8 event: PatientRegistered, SOAPSubmitted, DoctorNoteSubmitted, PrescriptionCreated, DepartmentAssigned, RecordUpdated, RoleAssigned, RecordDeactivated"),
    ("4",  "File .sol untuk disubmit",                 "✅ LULUS","contracts/EMRv2.sol — deployed ke Sepolia"),
    ("5",  "Laporan .docx dokumentasi fungsi",         "✅ LULUS","File ini — mencakup semua 18 fungsi, 4 modifier, 8 event"),
    ("6",  "Aplikasi yang berjalan",                   "✅ LULUS","https://emr-blockchain.vercel.app"),
    ("7",  "Smart contract deployed",                  "✅ LULUS","Sepolia: 0xBdc34e6C796235A1399f2E18e350e62E24b1Dca5"),
    ("8",  "Integrasi blockchain dengan web app",      "✅ LULUS","ethers.js v6 mengintegrasikan semua fungsi contract"),
    ("9",  "GitHub repository",                        "✅ LULUS","https://github.com/ervandyr2512/emr-blockchain"),
    ("10", "Semua 4 kategori CRUD terpenuhi",          "✅ LULUS","C: 7, R: 8, U: 2, D: 2"),
]:
    add_row(t, row, bold_first=True)

spacer()
subheading("B. Rekap Lengkap 18 Fungsi CRUD")
t2 = styled_table(["ID","Kategori","Nama Fungsi","Peran yang Dapat Memanggil"],
                  col_widths=[0.5,0.8,1.8,3.4])
cat_bg = {"CREATE":"E8F5E9","READ":"E3F2FD","UPDATE":"FFF8E1","DELETE":"FCE4EC"}
for row in [
    ("C1","CREATE","assignRole",          "Admin, Owner"),
    ("C2","CREATE","selfRegister",        "Semua wallet baru"),
    ("C3","CREATE","registerPatient",     "Semua peran terdaftar"),
    ("C4","CREATE","submitSOAP",          "Semua peran terdaftar"),
    ("C5","CREATE","submitDoctorNote",    "Semua peran terdaftar"),
    ("C6","CREATE","fulfillPrescription", "Semua peran terdaftar"),
    ("C7","CREATE","assignDepartment",    "Admin, Owner"),
    ("R1","READ",  "getEMRActions",       "Publik (view)"),
    ("R2","READ",  "getRole",             "Publik (view)"),
    ("R3","READ",  "isRegistered",        "Publik (view)"),
    ("R4","READ",  "getTotalActions",     "Publik (view)"),
    ("R5","READ",  "getUserProfile",      "Publik (view)"),
    ("R6","READ",  "getActionCount",      "Publik (view)"),
    ("R7","READ",  "getActionByIndex",    "Publik (view)"),
    ("R8","READ",  "getLatestAction",     "Publik (view)"),
    ("U1","UPDATE","updateRecord",        "Semua peran terdaftar"),
    ("U2","UPDATE","updateUserName",      "User sendiri / Admin / Owner"),
    ("D1","DELETE","deactivateRecord",    "Admin, Owner"),
    ("D2","DELETE","deactivateUser",      "Admin, Owner"),
]:
    add_row(t2, row, bold_first=True, bg=cat_bg.get(row[1]))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# XVI. KESIMPULAN
# ══════════════════════════════════════════════════════════════════════════════

heading("XVI. KESIMPULAN")
para("Proyek EMR Blockchain telah berhasil dibangun, di-deploy, dan diperbarui menjadi sistem rekam medis elektronik yang aman, transparan, dan dapat diaudit. Berikut pencapaian final:")
spacer()

for title, desc in [
    ("Smart Contract Komprehensif",
     "18 fungsi (C7+R8+U2+D2), 4 modifier, 8 event — melebihi semua persyaratan minimal tugas."),
    ("Alur Dokter Mandiri",
     "Dokter dapat mendaftar sendiri via /register-dokter dengan kredensial profesional lengkap (STR, SIP, spesialisasi). Admin mereview dan menyetujui/menolak dari dashboard."),
    ("Blockchain Trail Per Entri",
     "Setiap catatan SOAP dan pemeriksaan dokter memiliki riwayat lengkap txHash beserta informasi siapa yang membuat/memperbarui dan kapan."),
    ("Keamanan & Integritas",
     "SHA-256 hashing + blockchain immutability memastikan data tidak dapat dimanipulasi. Setiap perubahan meninggalkan jejak permanen."),
    ("Multi-Role Frontend",
     "6 peran dengan UI dan akses yang sepenuhnya terpisah. Sidebar, routing, dan fungsi hanya menampilkan yang relevan per peran."),
    ("Robustness",
     "safeFormat() utility mencegah crash 'Invalid Date' untuk semua data Firebase yang mungkin tidak memiliki timestamp lengkap."),
    ("Navigasi Konsisten",
     "Logo di semua halaman auth dapat diklik untuk kembali ke landing page. UX konsisten di seluruh aplikasi."),
]:
    bullet("", bold_prefix=f"{title}: ")
    p = doc.paragraphs[-1]
    set_run_font(p.add_run(desc), size=11)

spacer(2)
para("Smart contract EMRv2.sol merupakan file final yang di-deploy ke Ethereum Sepolia Testnet dan direkomendasikan untuk disubmit sebagai tugas.",
     italic=True, color=(0x1a,0x37,0x6c))

spacer(2); divider(); spacer()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("— Laporan Final Proyek EMR Blockchain —"), italic=True, size=10, color=(0x80,0x80,0x80))
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p2.add_run("Ervandy Rangganata  ·  2024 / 2025"), size=10, color=(0x80,0x80,0x80))

# ── Save ──────────────────────────────────────────────────────────────────────
output = "/Users/ervandyrangganata/Downloads/Blockchain - Smart Contract - EMR/Laporan_EMR_Blockchain_Ervandy_Rangganata.docx"
doc.save(output)
print(f"✅  Report saved: {output}")
